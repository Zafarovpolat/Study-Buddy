# backend/app/services/text_extractor.py
import os
import re
from pathlib import Path


def clean_text_for_db(text: str) -> str:
    """Очищает текст от символов, несовместимых с PostgreSQL UTF-8"""
    if not text:
        return ""
    
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = text.encode('utf-8', errors='replace').decode('utf-8')
    
    return text


class TextExtractor:
    """Извлечение текста из разных форматов файлов"""
    
    @staticmethod
    async def extract_from_pdf(file_path: str) -> str:
        """Извлечь текст из PDF, с fallback на OCR"""
        try:
            import pypdf
            
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            # Если текста нет или слишком мало — пробуем OCR
            if not text.strip() or len(text.strip()) < 50:
                print("📷 PDF без текста, пробуем OCR через Gemini...")
                try:
                    text = await TextExtractor._pdf_ocr_gemini(file_path)
                except Exception as ocr_error:
                    print(f"❌ OCR failed: {ocr_error}")
                    raise ValueError("PDF не содержит текста. Попробуйте загрузить как изображение.")
            
            return clean_text_for_db(text)
            
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Ошибка чтения PDF: {str(e)}")
    
    @staticmethod
    async def _pdf_ocr_gemini(file_path: str) -> str:
        """OCR для PDF через Gemini"""
        import google.generativeai as genai
        from app.core.config import settings
        import base64
        
        with open(file_path, 'rb') as f:
            pdf_data = f.read()
        
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel(settings.GEMINI_MODEL)
        
        response = model.generate_content([
            {
                "mime_type": "application/pdf",
                "data": base64.b64encode(pdf_data).decode('utf-8')
            },
            "Извлеки весь текст из этого PDF документа. Сохрани структуру. Только текст, без комментариев."
        ])
        
        text = response.text.strip()
        
        if not text or len(text) < 10:
            raise ValueError("Не удалось распознать текст")
        
        return text
    
    @staticmethod
    async def extract_from_docx(file_path: str) -> str:
        """Извлечь текст из DOCX"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            
            if not text.strip():
                raise ValueError("DOCX не содержит текста")
            
            return clean_text_for_db(text)
            
        except KeyError:
            raise ValueError("Файл повреждён. Сохраните как .docx в Word")
        except Exception as e:
            if "relationship" in str(e).lower():
                raise ValueError("Файл повреждён. Сохраните как .docx в Word")
            raise ValueError(f"Ошибка чтения DOCX: {str(e)}")
    
    @staticmethod
    async def extract_from_doc(file_path: str) -> str:
        """DOC не поддерживается"""
        raise ValueError(
            "Формат .doc не поддерживается. "
            "Откройте в Word → Файл → Сохранить как → .docx"
        )
    
    @staticmethod
    async def extract_from_txt(file_path: str) -> str:
        """Извлечь текст из TXT"""
        encodings = ['utf-8', 'cp1251', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    if text.strip():
                        return clean_text_for_db(text)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Не удалось прочитать файл")
    
    @staticmethod
    async def extract_from_image(file_path: str) -> str:
        """Извлечь текст из изображения через Gemini Vision"""
        try:
            import google.generativeai as genai
            from app.core.config import settings
            import base64
            
            genai.configure(api_key=settings.GEMINI_API_KEY)
            
            with open(file_path, 'rb') as f:
                image_data = f.read()
            
            ext = Path(file_path).suffix.lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.webp': 'image/webp',
                '.gif': 'image/gif'
            }
            mime_type = mime_types.get(ext, 'image/jpeg')
            
            model = genai.GenerativeModel(settings.GEMINI_MODEL)
            
            response = model.generate_content([
                {
                    "mime_type": mime_type,
                    "data": base64.b64encode(image_data).decode('utf-8')
                },
                "Извлеки весь текст с изображения. Сохрани структуру. Только текст, без комментариев."
            ])
            
            text = response.text.strip()
            
            if not text or len(text) < 3:
                raise ValueError("Текст не распознан")
            
            return clean_text_for_db(text)
            
        except Exception as e:
            error = str(e)
            if "404" in error or "not found" in error.lower():
                raise ValueError(f"Модель недоступна")
            raise ValueError(f"Ошибка OCR: {error[:100]}")
    
    @classmethod
    async def extract(cls, file_path: str, material_type: str) -> str:
        """Универсальный метод"""
        
        if not os.path.exists(file_path):
            raise ValueError("Файл не найден")
        
        if os.path.getsize(file_path) == 0:
            raise ValueError("Файл пустой")
        
        ext = Path(file_path).suffix.lower()
        
        extractors = {
            '.pdf': cls.extract_from_pdf,
            '.docx': cls.extract_from_docx,
            '.doc': cls.extract_from_doc,
            '.txt': cls.extract_from_txt,
            '.jpg': cls.extract_from_image,
            '.jpeg': cls.extract_from_image,
            '.png': cls.extract_from_image,
            '.webp': cls.extract_from_image,
        }
        
        extractor = extractors.get(ext)
        
        if not extractor:
            raise ValueError(f"Формат {ext} не поддерживается")
        
        print(f"📂 Extracting {ext} from {file_path}")
        
        text = await extractor(file_path)
        text = clean_text_for_db(text)
        text = re.sub(r'\n{3,}', '\n\n', text.strip())
        
        return text